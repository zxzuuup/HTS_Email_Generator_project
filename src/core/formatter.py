# src/core/formatter.py
from docx import Document
from docx.shared import RGBColor
from config.settings import GREETINGS, CLOSINGS, SIGNATURE


class WordFormatter:
    def __init__(self, file_name):
        self.word_doc = Document()
        self.file_name = file_name


    @staticmethod
    def _apply_word_style(run, style):
        """为 python-docx 的 run 应用样式"""
        if style == 'red':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
        elif style == 'bold':
            run.bold = True
        elif style == 'redbold':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
            run.bold = True
        # normal 样式无需特殊处理

    def add_paragraph(self, content=None, before_black=0, after_black=0):
        for i in range(before_black):
            self.word_doc.add_paragraph()

        if content:
            self.word_doc.add_paragraph(content)

        for i in range(after_black):
            self.word_doc.add_paragraph()

    def format_and_save_word(self, text_parts, style_parts, code, language):
        """格式化内容并保存为 Word 文档"""
        greeting = GREETINGS.get(language, "Hello Seller,")
        closing = CLOSINGS.get(language, "If you have any questions, please contact us in time. Thanks！")

        self.word_doc.add_paragraph(greeting)
        self.word_doc.add_paragraph()

        for i, (text, styles) in enumerate(zip(text_parts, style_parts)):
            if i > 0:
                self.word_doc.add_paragraph()

            question_title = f"Question {i + 1}:" if language == 'EN' else f"问题 {i + 1}:"
            title_para = self.word_doc.add_paragraph()
            title_para.add_run(question_title)

            content_para = self.word_doc.add_paragraph()
            full_text = text

            if not styles:
                content_para.add_run(full_text)
            else:
                last_end_idx = 0
                for part_text, part_style in styles:
                    start_idx = full_text.find(part_text, last_end_idx)
                    if start_idx == -1:
                        run = content_para.add_run(part_text)
                        WordFormatter._apply_word_style(run, 'normal')
                        continue

                    if start_idx > last_end_idx:
                        prefix = full_text[last_end_idx:start_idx]
                        run = content_para.add_run(prefix)
                        WordFormatter._apply_word_style(run, 'normal')

                    styled_run = content_para.add_run(part_text)
                    WordFormatter._apply_word_style(styled_run, part_style)

                    last_end_idx = start_idx + len(part_text)

                if last_end_idx < len(full_text):
                    suffix = full_text[last_end_idx:]
                    run = content_para.add_run(suffix)
                    WordFormatter._apply_word_style(run, 'normal')

        self.word_doc.add_paragraph()
        self.word_doc.add_paragraph(closing)


    def save(self):
        self.word_doc.add_paragraph()
        self.word_doc.add_paragraph(SIGNATURE)
        self.word_doc.save(self.file_name)
